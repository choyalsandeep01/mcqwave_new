import os
import django
import traceback
import logging
import docx
import re
from datetime import datetime

# Set up Django environment
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'myproject.settings')
django.setup()

# Import your models
from mcqs.models import Subject, Unit, Chapter, Topic, difficulties, mcq_types, MCQ

# Configure logging
logging.basicConfig(level=logging.DEBUG, 
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    filename='mcq_upload_log.txt')
logger = logging.getLogger(__name__)

def normalize_text(text):
    """
    Normalize text to match database entries
    Converts to title case, ensuring first letter is capital
    """
    return text.strip().title()

def validate_difficulty(difficulty_name):
    """
    Validate if difficulty exists in the database
    """
    try:
        normalized_difficulty = normalize_text(difficulty_name)
        difficulty = difficulties.objects.filter(name=normalized_difficulty).first()
        
        if not difficulty:
            logger.error(f"Difficulty '{difficulty_name}' does not exist in the database.")
            return None
        
        return difficulty
    except Exception as e:
        logger.error(f"Error validating difficulty {difficulty_name}: {e}")
        return None

def validate_mcq_type(mcq_type):
    """
    Validate if MCQ type exists in the database
    """
    try:
        normalized_type = normalize_text(mcq_type)
        mcq_type_obj = mcq_types.objects.filter(types=normalized_type).first()
        
        if not mcq_type_obj:
            logger.error(f"MCQ Type '{mcq_type}' does not exist in the database.")
            return None
        
        return mcq_type_obj
    except Exception as e:
        logger.error(f"Error validating MCQ type {mcq_type}: {e}")
        return None

def parse_mcq(paragraphs):
    """
    Parse MCQ entry with enhanced flexibility
    Handles both single-line and multi-line formats
    """
    # Join all paragraphs, removing extra whitespaces
    full_text = ' '.join([p.text.strip() for p in paragraphs])
    full_text = re.sub(r'\s+', ' ', full_text)
    
    logger.debug(f"Full MCQ text: {full_text}")
    
    try:
        # First try to split by '|'
        parts = [p.strip() for p in full_text.split('|')]
        
        # Validate minimum number of parts
        if len(parts) < 7:
            # If splitting fails, try alternative parsing
            logger.error(f"Insufficient parts in MCQ: {full_text}")
            return None
        
        # Extract question (first part)
        question = parts[0]
        
        # Handling options and correct answer
        if len(parts) >= 7:
            options = parts[1:5]
            correct_answer = parts[5]
            explanation = parts[6]
            
            # Default difficulty and type if not specified
            difficulty = parts[7] if len(parts) > 7 else 'Easy'
            mcq_type = parts[8] if len(parts) > 8 else 'General'
            
            # Construct full parts list
            full_parts = [
                question,      # 0: Question
                options[0],    # 1: Option 1
                options[1],    # 2: Option 2
                options[2],    # 3: Option 3
                options[3],    # 4: Option 4
                correct_answer,# 5: Correct Option
                explanation,   # 6: Explanation
                difficulty,    # 7: Difficulty
                mcq_type      # 8: MCQ Type
            ]
            
            logger.debug(f"Parsed MCQ parts: {full_parts}")
            return full_parts
        
        return None
    
    except Exception as e:
        logger.error(f"Error parsing MCQ: {e}")
        logger.error(f"Full text: {full_text}")
        return None

def process_docx_file(file_path):
    """
    Process DOCX file and upload to database with better section handling
    """
    try:
        doc = docx.Document(file_path)
        subject_name = os.path.splitext(os.path.basename(file_path))[0]
        subject, _ = Subject.objects.get_or_create(name=subject_name)
        
        current_unit = None
        current_chapter = None
        current_topic = None
        uploaded_content = []
        uploaded_mcq_indices = []
        
        i = 0
        while i < len(doc.paragraphs):
            text = doc.paragraphs[i].text.strip()
            
            if not text:
                i += 1
                continue
                
            # Handle section headers
            if 'Unit-' in text:
                unit_name = text.split('Unit-', 1)[1].strip()
                current_unit, _ = Unit.objects.get_or_create(
                    subject=subject,
                    name=unit_name
                )
                current_chapter = None
                current_topic = None
                i += 1
                continue
                
            elif 'Chapter-' in text:
                if not current_unit:
                    current_unit, _ = Unit.objects.get_or_create(
                        subject=subject,
                        name="Default Unit"
                    )
                chapter_name = text.split('Chapter-', 1)[1].strip()
                current_chapter, _ = Chapter.objects.get_or_create(
                    unit=current_unit,
                    name=chapter_name
                )
                current_topic = None
                i += 1
                continue
                
            elif 'Topic-' in text:
                if not current_chapter:
                    if not current_unit:
                        current_unit, _ = Unit.objects.get_or_create(
                            subject=subject,
                            name="Default Unit"
                        )
                    current_chapter, _ = Chapter.objects.get_or_create(
                        unit=current_unit,
                        name="Default Chapter"
                    )
                topic_name = text.split('Topic-', 1)[1].strip()
                current_topic, _ = Topic.objects.get_or_create(
                    chapter=current_chapter,
                    name=topic_name
                )
                i += 1
                continue
            
            # MCQ Processing
            if '|' in text:
                mcq_text = text
                mcq_start_index = i
                
                # Collect multi-line MCQ if necessary
                while i + 1 < len(doc.paragraphs):
                    next_text = doc.paragraphs[i + 1].text.strip()
                    if not next_text or any(header in next_text for header in ['Unit-', 'Chapter-', 'Topic-']):
                        break
                    if '|' in next_text:  # Start of new MCQ
                        break
                    mcq_text += ' ' + next_text
                    i += 1
                
                # Process the MCQ
                parts = parse_mcq([docx.Document().add_paragraph(mcq_text)])
                if parts and len(parts) >= 8:
                    try:
                        difficulty = validate_difficulty(parts[7])
                        mcq_type = validate_mcq_type(parts[8] if len(parts) > 8 else 'General')
                        
                        if all([current_chapter, current_topic, difficulty, mcq_type]):
                            mcq = MCQ.objects.create(
                                text=parts[0],
                                option_1=parts[1],
                                option_2=parts[2],
                                option_3=parts[3],
                                option_4=parts[4],
                                correct_option=parts[5],
                                explanation=parts[6],
                                topic=current_topic,
                                difficulty=difficulty,
                                types=mcq_type
                            )
                            
                            current_context = {
                                'unit': f"Unit-{current_unit.name}" if current_unit else None,
                                'chapter': f"Chapter-{current_chapter.name}" if current_chapter else None,
                                'topic': f"Topic-{current_topic.name}" if current_topic else None
                            }
                            
                            uploaded_content.append({
                                'context': current_context,
                                'paragraphs': [mcq_text]
                            })
                            
                            uploaded_mcq_indices.append(mcq_start_index)
                            logger.info(f"Successfully uploaded MCQ: {mcq.text[:50]}...")
                    except Exception as e:
                        logger.error(f"Error creating MCQ: {e}")
                        logger.error(traceback.format_exc())
            
            i += 1
        
        # Create output documents
        create_output_documents(doc, file_path, uploaded_mcq_indices, uploaded_content, subject_name)
        
    except Exception as file_error:
        logger.error(f"Error processing file {file_path}: {file_error}")
        logger.error(traceback.format_exc())
def process_buffered_mcq(mcq_buffer, current_unit, current_chapter, current_topic, 
                        uploaded_content, uploaded_mcq_indices, buffer_start_index):
    """Helper function to process MCQ buffer and create database entry"""
    # Ensure proper hierarchy exists
    if not current_chapter:
        logger.warning("MCQ found without chapter context. Skipping.")
        return
    
    if not current_topic:
        current_topic, _ = Topic.objects.get_or_create(
            chapter=current_chapter,
            name=current_chapter.name
        )
    
    # Parse MCQ
    parts = parse_mcq(mcq_buffer)
    if not parts or len(parts) < 8:
        return
    
    # Validate difficulty and type
    difficulty = validate_difficulty(parts[7])
    mcq_type = validate_mcq_type(parts[8] if len(parts) > 8 else 'General')
    
    if not all([difficulty, mcq_type]):
        return
    
    try:
        # Create MCQ
        mcq = MCQ.objects.create(
            text=parts[0],
            option_1=parts[1],
            option_2=parts[2],
            option_3=parts[3],
            option_4=parts[4],
            correct_option=parts[5],
            explanation=parts[6],
            topic=current_topic,
            difficulty=difficulty,
            types=mcq_type
        )
        
        # Track successful upload
        current_context = {
            'unit': f"Unit-{current_unit.name}" if current_unit else None,
            'chapter': f"Chapter-{current_chapter.name}" if current_chapter else None,
            'topic': f"Topic-{current_topic.name}" if current_topic else None
        }
        
        uploaded_content.append({
            'context': current_context,
            'paragraphs': [p.text for p in mcq_buffer]
        })
        
        uploaded_mcq_indices.extend(range(buffer_start_index, buffer_start_index + len(mcq_buffer)))
        
        logger.info(f"Successfully uploaded MCQ: {mcq.text[:50]}...")
    
    except Exception as e:
        logger.error(f"Error creating MCQ: {e}")
        logger.error(traceback.format_exc())

def main():
    """
    Main function to process MCQ files
    """
    mcq_directory = os.getcwd()
    # Directory containing MCQ DOCX files
    
    # Process all DOCX files in the directory
    for filename in os.listdir(mcq_directory):
        if filename.endswith('.docx'):
            file_path = os.path.join(mcq_directory, filename)
            process_docx_file(file_path)

def create_output_documents(doc, file_path, uploaded_mcq_indices, uploaded_content, subject_name):
    """Helper function to create output documents"""
    # Create new document with remaining content
    remaining_doc = docx.Document()
    for i, para in enumerate(doc.paragraphs):
        if i not in uploaded_mcq_indices:
            remaining_doc.add_paragraph(para.text)
    
    # Save updated original document
    remaining_doc.save(file_path)
    
    # Create document with uploaded content
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    uploaded_doc = docx.Document()
    uploaded_doc.add_heading(f'Uploaded MCQs - {subject_name}', 0)
    
    # Track current context
    current_unit = None
    current_chapter = None
    current_topic = None
    
    # Add uploaded content with context
    for content in uploaded_content:
        if content['context']['unit'] != current_unit:
            uploaded_doc.add_paragraph(content['context']['unit'])
            current_unit = content['context']['unit']
        
        if content['context']['chapter'] != current_chapter:
            uploaded_doc.add_paragraph(content['context']['chapter'])
            current_chapter = content['context']['chapter']
        
        if content['context']['topic'] != current_topic:
            uploaded_doc.add_paragraph(content['context']['topic'])
            current_topic = content['context']['topic']
        
        for para_text in content['paragraphs']:
            uploaded_doc.add_paragraph(para_text)
        
        uploaded_doc.add_paragraph()
    
    # Save uploaded content document
    uploaded_file_path = os.path.join(
        os.path.dirname(file_path),
        f'uploaded_{os.path.basename(file_path).replace(".docx", "")}_{timestamp}.docx'
    )
    uploaded_doc.save(uploaded_file_path)
    
if __name__ == '__main__':
    # Install required libraries
    try:
        import docx
    except ImportError:
        print("Installing required libraries...")
        import subprocess
        subprocess.check_call(['pip', 'install', 'python-docx', 'django'])
    
    main()
    print("MCQ upload process completed. Check mcq_upload_log.txt for details.")